"""
DocFlow — DOCX Parser
Extracts text, headings, tables, images, and metadata from .docx files.
"""
from __future__ import annotations

from pathlib import Path
from typing import Any

from loguru import logger

from utils.markdown_utils import tables_to_markdown, clean_markdown


class DOCXParser:
    async def parse(self, file_path: Path, options: Any) -> dict:
        logger.info(f"[DOCXParser] Parsing: {file_path.name}")
        try:
            from docx import Document
            from docx.shared import Pt
        except ImportError:
            return self._empty_result("python-docx not installed")

        doc = Document(str(file_path))
        md_parts: list[str] = []
        tables_data: list[dict] = []
        images: list[dict] = []

        # Core properties / metadata
        props = doc.core_properties
        metadata = {
            "title": props.title or "",
            "author": props.author or "",
            "subject": props.subject or "",
            "keywords": props.keywords or "",
            "created": str(props.created or ""),
            "modified": str(props.modified or ""),
            "revision": props.revision or 0,
        }

        # Paragraphs
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                md_parts.append("")
                continue

            style_name = para.style.name.lower() if para.style else ""
            if "heading 1" in style_name:
                md_parts.append(f"\n# {text}\n")
            elif "heading 2" in style_name:
                md_parts.append(f"\n## {text}\n")
            elif "heading 3" in style_name:
                md_parts.append(f"\n### {text}\n")
            elif "heading" in style_name:
                md_parts.append(f"\n#### {text}\n")
            elif "list" in style_name:
                md_parts.append(f"- {text}")
            else:
                # Inline formatting
                para_md = ""
                for run in para.runs:
                    run_text = run.text
                    if run.bold and run.italic:
                        run_text = f"***{run_text}***"
                    elif run.bold:
                        run_text = f"**{run_text}**"
                    elif run.italic:
                        run_text = f"*{run_text}*"
                    para_md += run_text
                md_parts.append(para_md)

        # Tables
        if options.extract_tables:
            for i, table in enumerate(doc.tables):
                rows: list[list[str]] = []
                for row in table.rows:
                    rows.append([cell.text.strip() for cell in row.cells])
                if rows:
                    md = tables_to_markdown(rows)
                    md_parts.append(f"\n{md}\n")
                    tables_data.append({"index": i, "data": rows, "markdown": md})

        # Inline images count
        if options.extract_images:
            for rel in doc.part.rels.values():
                if "image" in rel.reltype:
                    images.append({"rel_id": rel.relid, "target": str(rel.target_ref)})

        full_md = "\n".join(md_parts)
        return {
            "markdown": clean_markdown(full_md),
            "text": "\n".join(p.lstrip("#*- ") for p in md_parts if p.strip()),
            "structured": {"paragraph_count": len(doc.paragraphs), "table_count": len(doc.tables)},
            "tables": tables_data,
            "images": images,
            "metadata": metadata,
            "hyperlinks": [],
        }

    @staticmethod
    def _empty_result(reason: str) -> dict:
        return {"markdown": f"<!-- Error: {reason} -->", "text": "", "structured": {},
                "tables": [], "images": [], "metadata": {}, "hyperlinks": []}
